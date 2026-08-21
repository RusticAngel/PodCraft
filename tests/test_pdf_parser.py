import os
import sys
import tempfile

import pytest

sys.path.insert(0, os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

from src.phase2_document_processing.pdf_parser import PDFScriptParser
from src.phase2_document_processing.script_analyzer import ScriptAnalyzer
from src.phase2_document_processing.speaker_identifier import SpeakerIdentifier
from scripts.make_demo_pdf import build_pdf


@pytest.fixture(scope="module")
def sample_pdf():
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(build_pdf([
            "PODCAST: TEST",
            "INT. STUDIO",
            "",
            "HOST: Hello and welcome to the show",
            "GUEST: Thanks for having me, this is amazing",
            "HOST: Let's talk about AI today",
        ]))
        path = tmp.name
    yield path
    os.unlink(path)


def test_parse_returns_expected_keys(sample_pdf):
    data = PDFScriptParser().parse(sample_pdf)
    assert set(["full_text", "speakers", "dialogue_segments", "topics", "mood", "scene_breaks", "estimated_duration"]).issubset(data.keys())
    assert "HOST" in [s.upper() for s in data["speakers"]]
    assert data["scene_breaks"] >= 1


def test_dialogue_segments(sample_pdf):
    data = PDFScriptParser().parse(sample_pdf)
    speakers = [seg["speaker"] for seg in data["dialogue_segments"]]
    assert "HOST" in speakers
    assert all(seg["text"] for seg in data["dialogue_segments"])


def test_duration_estimate_positive(sample_pdf):
    data = PDFScriptParser().parse(sample_pdf)
    assert data["estimated_duration"] >= 0


def test_script_analyzer(sample_pdf):
    data = PDFScriptParser().parse(sample_pdf)
    analysis = ScriptAnalyzer().analyze(data)
    assert "structure" in analysis
    assert "production_cues" in analysis
    assert analysis["segment_count"] == len(data["dialogue_segments"])


def test_speaker_identifier_assigns_voices(sample_pdf):
    data = PDFScriptParser().parse(sample_pdf)
    profiles = SpeakerIdentifier().identify(data["speakers"], data["dialogue_segments"])
    assert profiles
    assert all(p["voice"] and p["role"] for p in profiles)
    # Two profiles for two speakers
    assert len(profiles) == len(data["speakers"])


def test_speaker_identifier_distinct_voices():
    identifier = SpeakerIdentifier()
    profiles = [
        {"speaker": "a", "voice": "v1"},
        {"speaker": "b", "voice": "v2"},
    ]
    assert identifier.assign_voice("A", profiles) == "v1"
    assert identifier.assign_voice("b", profiles) == "v2"
    assert identifier.assign_voice("unknown", profiles) == identifier.default_voice


def test_speaker_role_inference_parenthesized_and_support():
    """'OUTRO (Narrator)' must be inferred as a support role, not host:
    the parenthesized label must not leak into hint matching."""
    identifier = SpeakerIdentifier()
    assert identifier._infer_role("OUTRO (Narrator)") == "support"
    assert identifier._infer_role("Intro (Producer)") == "support"
    assert identifier._infer_role("HOST") == "host"
    assert identifier._infer_role("Narrator") == "host"
    assert identifier._infer_role("GUEST") == "guest"


def test_parser_skips_metadata_and_handles_parenthesized_speakers():
    """FORMAT:/TOPIC: headers must not become speakers; 'OUTRO (Narrator):'
    with a parenthetical must be recognized as a speaker."""
    with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as tmp:
        tmp.write(build_pdf([
            "PODCAST TITLE: THE CREATOR'S CUT",
            "FORMAT: INTERVIEW",
            "DURATION: 3 MINUTES",
            "INT. STUDIO - MORNING",
            "",
            "HOST: Welcome back!",
            "GUEST: Thanks for having me.",
            "HOST: Let's dig in.",
            "OUTRO (Narrator): Thanks for listening!",
        ]))
        path = tmp.name
    try:
        data = PDFScriptParser().parse(path)
    finally:
        os.unlink(path)

    speakers = data["speakers"]
    assert "format" not in speakers
    assert "duration" not in speakers
    assert "outro (narrator)" in speakers

    seg_speakers = [seg["speaker"] for seg in data["dialogue_segments"]]
    assert seg_speakers == ["HOST", "GUEST", "HOST", "OUTRO (Narrator)"]


_SCRIPT_LINES = [
    "HOST: Hello and welcome to the show",
    "GUEST: Thanks for having me, this is amazing",
    "OUTRO (Narrator): See you next week!",
]


def test_parser_supports_txt_md_docx(tmp_path):
    """The same script must parse identically from TXT, MD and DOCX."""
    import docx as docx_lib

    txt = tmp_path / "script.txt"
    txt.write_text("\n".join(_SCRIPT_LINES), encoding="utf-8")

    md = tmp_path / "script.md"
    md.write_text("# Episode 1\n\n" + "\n".join(_SCRIPT_LINES), encoding="utf-8")

    docx_path = tmp_path / "script.docx"
    document = docx_lib.Document()
    document.add_paragraph("Episode 1")
    for line in _SCRIPT_LINES:
        document.add_paragraph(line)
    document.save(str(docx_path))

    expected_speakers = ["host", "guest", "outro (narrator)"]
    expected_segments = ["HOST", "GUEST", "OUTRO (Narrator)"]
    for path in (txt, md, docx_path):
        data = PDFScriptParser().parse(str(path))
        assert data["speakers"] == expected_speakers, f"speakers differ for {path.suffix}"
        got = [s["speaker"] for s in data["dialogue_segments"]]
        assert got == expected_segments, f"segments differ for {path.suffix}"
        assert all(s["text"] for s in data["dialogue_segments"])


def test_parser_rejects_unsupported_extension(tmp_path):
    bad = tmp_path / "script.doc"
    bad.write_bytes(b"old binary word doc")
    with pytest.raises(ValueError, match="Unsupported script format"):
        PDFScriptParser().parse(str(bad))


def test_supported_extensions_listed():
    exts = PDFScriptParser.supported_extensions()
    assert exts == [".docx", ".markdown", ".md", ".pdf", ".txt"]


_NUMBERED_SCRIPT = """Speaker 1 – Host:
Welcome back to The Final Whistle. Today we're looking at one of the biggest questions in sport.

Speaker 2 – Analyst:
Talent helps, of course. But talent alone doesn't win championships.

Speaker 3 – Former Player:
I learned that the hard way. You can have the fastest players in the league.

Speaker 1 – Host:
And that's what makes sport so unpredictable.

Speaker 2 – Analyst:
Exactly. Statistics can tell us who should win, but they can't measure determination.

Speaker 3 – Former Player:
And sometimes, that's all you need.

Speaker 1 – Host:
Until next time, keep playing, keep believing.
"""


def _write_numbered_script(path, dash="–"):
    body = _NUMBERED_SCRIPT.replace("–", dash)
    path.write_text(body, encoding="utf-8")
    return str(path)


def test_parser_supports_numbered_speaker_labels(tmp_path):
    """'Speaker N – Role:' style labels must be recognized as speakers."""
    data = PDFScriptParser().parse(_write_numbered_script(tmp_path / "numbered.txt"))
    assert data["speakers"] == [
        "speaker 1 – host",
        "speaker 2 – analyst",
        "speaker 3 – former player",
    ]
    segments = data["dialogue_segments"]
    assert len(segments) == 7
    assert [s["speaker"] for s in segments] == [
        "Speaker 1 – Host", "Speaker 2 – Analyst", "Speaker 3 – Former Player",
        "Speaker 1 – Host", "Speaker 2 – Analyst", "Speaker 3 – Former Player",
        "Speaker 1 – Host",
    ]
    assert all(s["text"] for s in segments)
    # Label must not leak into the dialogue text.
    assert not segments[0]["text"].lower().startswith("speaker")


def test_parser_numbered_labels_hyphen_and_emdash(tmp_path):
    """Hyphen and em-dash variants parse identically."""
    for name, dash in (("hyphen.txt", "-"), ("emdash.txt", "—")):
        data = PDFScriptParser().parse(_write_numbered_script(tmp_path / name, dash=dash))
        assert len(data["speakers"]) == 3
        assert len(data["dialogue_segments"]) == 7


def test_numbered_speaker_role_inference():
    """Role hints inside numbered labels still drive role assignment."""
    identifier = SpeakerIdentifier()
    assert identifier._infer_role("Speaker 1 – Host") == "host"
    assert identifier._infer_role("Speaker 2 - Analyst") == "guest"
    assert identifier._infer_role("Speaker 3 – Producer") == "support"


def test_numbered_labels_do_not_break_metadata_filtering(tmp_path):
    """Metadata headers stay filtered when both patterns are active."""
    body = ("FORMAT: INTERVIEW\n\n" + _NUMBERED_SCRIPT)
    path = tmp_path / "mixed.txt"
    path.write_text(body, encoding="utf-8")
    data = PDFScriptParser().parse(str(path))
    assert "format" not in data["speakers"]
    assert len(data["dialogue_segments"]) == 7